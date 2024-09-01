import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Asistente de Investigación", page_icon="🔍", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Asistente de Investigación que ayuda a los usuarios a encontrar y resumir artículos académicos relevantes sobre un tema específico.

    ### Cómo usar la aplicación:

    1. Ingrese un tema de investigación en el campo de texto.
    2. Haga clic en "Buscar artículos" para iniciar la búsqueda.
    3. Revise los resúmenes y puntos clave generados para cada artículo.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **[Tu Nombre]**, [Fecha Actual]

    ### Cómo citar esta aplicación (formato APA):
    [Tu Apellido], [Inicial del Nombre]. (2024). *Asistente de Investigación* [Aplicación web]. [URL de tu aplicación]

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar resúmenes y extraer puntos clave. Siempre verifique la información con las fuentes originales para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Asistente de Investigación")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    def buscar_articulos(query):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} site:scholar.google.com",
            "num": 10
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_resumen(titulo, snippet):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Título del artículo: {titulo}\n\nFragment del artículo: {snippet}\n\nGenera un resumen conciso del artículo basado en la información proporcionada. Luego, extrae y enumera 3 puntos clave del artículo.\n\nResumen:\n\nPuntos clave:\n1.",
            "max_tokens": 2048,
            "temperature": 0.2,
            "top_p": 0.9,
            "top_k": 50,
            "repetition_penalty": 1.1,
            "stop": ["Título del artículo:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(tema, resultados):
        doc = Document()
        doc.add_heading('Resultados de la Investigación', 0)

        doc.add_heading('Tema', level=1)
        doc.add_paragraph(tema)

        for resultado in resultados:
            doc.add_heading(resultado['title'], level=2)
            doc.add_paragraph(f"URL: {resultado['link']}")
            doc.add_paragraph(resultado['resumen'])

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con las fuentes originales para un análisis más profundo.')

        return doc

    tema_investigacion = st.text_input("Ingresa tu tema de investigación:")

    if st.button("Buscar artículos"):
        if tema_investigacion:
            with st.spinner("Buscando artículos y generando resúmenes..."):
                resultados_busqueda = buscar_articulos(tema_investigacion)
                resultados = []

                for item in resultados_busqueda.get("organic", []):
                    titulo = item.get("title", "")
                    snippet = item.get("snippet", "")
                    link = item.get("link", "")

                    resumen_y_puntos = generar_resumen(titulo, snippet)

                    resultado = {
                        "title": titulo,
                        "link": link,
                        "resumen": resumen_y_puntos
                    }
                    resultados.append(resultado)

                # Mostrar los resultados
                st.subheader(f"Resultados para el tema: {tema_investigacion}")
                for resultado in resultados:
                    st.markdown(f"### [{resultado['title']}]({resultado['link']})")
                    st.markdown(resultado['resumen'])
                    st.markdown("---")

                # Botón para descargar el documento
                doc = create_docx(tema_investigacion, resultados)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar resumen en DOCX",
                    data=buffer,
                    file_name=f"Investigacion_{tema_investigacion.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa un tema de investigación.")
